"""
BiblIA DOCX Export Module
==========================
Export transcriptions to DOCX with:
- RTL support (Hebrew/Arabic)
- Formatting preservation
- Paragraph styles
- Document metadata
"""

import io
from typing import Optional
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from arabic_reshaper import reshape
from bidi.algorithm import get_display

from django.conf import settings
from core.models import Document, DocumentPart, Line, LineTranscription, Transcription


class DOCXExporter:
    """
    Export transcriptions to DOCX format
    """
    
    def __init__(self, document: Document, transcription: Transcription):
        self.document = document
        self.transcription = transcription
        self.docx = DocxDocument()
        self.buffer = io.BytesIO()
        
        # Set up RTL support
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up document styles for RTL support"""
        # Main text style
        styles = self.docx.styles
        
        # Hebrew/Arabic paragraph style
        try:
            rtl_style = styles.add_style('RTL Text', WD_STYLE_TYPE.PARAGRAPH)
            rtl_style.font.name = 'David'  # Hebrew font
            rtl_style.font.size = Pt(12)
            rtl_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rtl_style.paragraph_format.line_spacing = 1.5
        except:
            # Style might already exist
            pass
        
        # LTR paragraph style
        try:
            ltr_style = styles.add_style('LTR Text', WD_STYLE_TYPE.PARAGRAPH)
            ltr_style.font.name = 'Calibri'
            ltr_style.font.size = Pt(12)
            ltr_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ltr_style.paragraph_format.line_spacing = 1.5
        except:
            pass
    
    def export(
        self,
        include_metadata: bool = True,
        include_images: bool = False,
        paragraph_per_line: bool = True
    ) -> io.BytesIO:
        """
        Export to DOCX with specified options
        
        Args:
            include_metadata: Include document metadata
            include_images: Include original images
            paragraph_per_line: Create a paragraph for each line (vs continuous text)
            
        Returns:
            BytesIO buffer with DOCX content
        """
        # Add title
        title = self.docx.add_heading(self.document.name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if include_metadata:
            self._add_metadata()
        
        # Process parts
        parts = DocumentPart.objects.filter(document=self.document).order_by('order')
        
        for part_num, part in enumerate(parts, 1):
            # Add part heading
            heading = self.docx.add_heading(f'Part {part_num}', level=2)
            
            # Add image if requested
            if include_images and part.image:
                try:
                    self.docx.add_picture(part.image.path, width=Inches(6))
                    last_paragraph = self.docx.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass
            
            # Add transcription text
            lines = Line.objects.filter(document_part=part).order_by('order')
            
            if paragraph_per_line:
                # Each line as separate paragraph
                for line in lines:
                    self._add_line_as_paragraph(line)
            else:
                # Continuous text
                self._add_lines_as_continuous_text(lines)
            
            # Add page break between parts
            if part_num < parts.count():
                self.docx.add_page_break()
        
        # Save to buffer
        self.docx.save(self.buffer)
        self.buffer.seek(0)
        return self.buffer
    
    def _add_metadata(self):
        """Add document metadata"""
        props = self.docx.core_properties
        props.title = self.document.name
        props.author = self.document.owner.username if self.document.owner else "BiblIA"
        props.subject = f"Transcription: {self.transcription.name}"
        props.comments = f"Created: {self.document.created_at}"
        props.category = "OCR Transcription"
        props.keywords = "BiblIA, eScriptorium, OCR, Transcription"
        
        # Add metadata paragraph
        meta = self.docx.add_paragraph()
        meta.add_run(f"Document: ").bold = True
        meta.add_run(f"{self.document.name}\n")
        meta.add_run(f"Transcription: ").bold = True
        meta.add_run(f"{self.transcription.name}\n")
        meta.add_run(f"Owner: ").bold = True
        meta.add_run(f"{self.document.owner.username if self.document.owner else 'N/A'}\n")
        meta.add_run(f"Parts: ").bold = True
        meta.add_run(f"{DocumentPart.objects.filter(document=self.document).count()}\n")
        meta.add_run(f"Lines: ").bold = True
        meta.add_run(f"{Line.objects.filter(document_part__document=self.document).count()}\n")
        meta.style = 'Normal'
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add separator
        self.docx.add_paragraph("_" * 50)
        self.docx.add_paragraph()
    
    def _format_rtl_text(self, text: str) -> str:
        """Format RTL text for DOCX"""
        if not text:
            return ""
        
        # Check if text contains Hebrew or Arabic
        has_hebrew = any('\u0590' <= char <= '\u05FF' for char in text)
        has_arabic = any('\u0600' <= char <= '\u06FF' for char in text)
        
        if has_hebrew or has_arabic:
            # Reshape Arabic if needed
            if has_arabic:
                text = reshape(text)
            # Apply BiDi algorithm
            return get_display(text)
        
        return text
    
    def _is_rtl_text(self, text: str) -> bool:
        """Check if text is RTL"""
        return any('\u0590' <= char <= '\u06FF' for char in text)
    
    def _add_line_as_paragraph(self, line: Line):
        """Add a single line as a paragraph"""
        try:
            lt = LineTranscription.objects.get(
                line=line,
                transcription=self.transcription
            )
            text = lt.content or ""
        except LineTranscription.DoesNotExist:
            text = ""
        
        if not text:
            return
        
        # Create paragraph
        p = self.docx.add_paragraph()
        
        # Determine if RTL
        is_rtl = self._is_rtl_text(text)
        
        # Format text
        formatted_text = self._format_rtl_text(text) if is_rtl else text
        
        # Add text with appropriate style
        run = p.add_run(formatted_text)
        
        # Set font based on script
        if is_rtl:
            run.font.name = 'David'  # Hebrew font
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_to_left = True
        else:
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run.font.size = Pt(12)
    
    def _add_lines_as_continuous_text(self, lines):
        """Add lines as continuous text in one paragraph"""
        p = self.docx.add_paragraph()
        
        for line in lines:
            try:
                lt = LineTranscription.objects.get(
                    line=line,
                    transcription=self.transcription
                )
                text = lt.content or ""
            except LineTranscription.DoesNotExist:
                continue
            
            if not text:
                continue
            
            # Determine if RTL
            is_rtl = self._is_rtl_text(text)
            
            # Format text
            formatted_text = self._format_rtl_text(text) if is_rtl else text
            
            # Add text
            run = p.add_run(formatted_text + " ")
            
            # Set font
            if is_rtl:
                run.font.name = 'David'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True
            else:
                run.font.name = 'Calibri'
            
            run.font.size = Pt(12)
